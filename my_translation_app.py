import tkinter as tk
from tkinter import filedialog, messagebox, Menu
from tkinter.ttk import Progressbar, Scale
import pyperclip
import threading
import time
import pdfplumber
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch
from langdetect import detect
import pystray
from PIL import Image, ImageDraw

# Load NLLB-200 Distilled 600M model
model_name = "facebook/nllb-200-distilled-600M"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
model.to(device)
model.eval()

# Enable mixed precision inference
if torch.cuda.is_available():
    model.half()  # Use FP16 for acceleration

# Supported languages with NLLB-200 language codes
languages = {
    "自动检测": "auto",
    "英语": "eng_Latn",
    "中文": "zho_Hans",
    "法语": "fra_Latn",
    "德语": "deu_Latn",
    "西班牙语": "spa_Latn",
    "日语": "jpn_Jpan",
    "韩语": "kor_Hang",
    "俄语": "rus_Cyrl"
}


# Translate single or batch text (supports GPU and batch processing)
def translate_text(texts, src_lang="auto", tgt_lang="zho_Hans"):
    if isinstance(texts, str):
        texts = [texts]
    if not any(text.strip() for text in texts):
        return [""] if len(texts) == 1 else ["" for _ in texts]

    # Auto-detect source language
    if src_lang == "auto":
        detected_lang = detect(texts[0])
        src_lang = next((code for name, code in languages.items() if detected_lang in code.lower()), "eng_Latn")

    # Tokenize input
    inputs = tokenizer(texts, return_tensors="pt", padding=True, truncation=True).to(device)

    # Get target language ID from vocabulary
    vocab = tokenizer.get_vocab()
    if tgt_lang not in vocab:
        print(f"Warning: Target language '{tgt_lang}' not in tokenizer vocabulary, defaulting to 'eng_Latn'")
        tgt_lang = "eng_Latn"
    tgt_lang_id = tokenizer.convert_tokens_to_ids(tgt_lang)

    # Generate translation
    with torch.no_grad():
        translated = model.generate(**inputs, forced_bos_token_id=tgt_lang_id)
    return [tokenizer.decode(t, skip_special_tokens=True) for t in translated]


# Extract text and layout from PDF
def extract_pdf_content(file_path):
    content = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            for char in page.chars:
                content.append({
                    "text": char["text"],
                    "x": char["x0"],
                    "y": A4[1] - char["y0"],
                    "size": char["size"],
                    "page": page_num
                })
    return content


# Generate translated PDF
def translate_pdf_document(file_path, src_lang, tgt_lang, progress_callback=None, stop_flag=None):
    content = extract_pdf_content(file_path)
    translated_file = file_path.rsplit(".", 1)[0] + "_translated.pdf"
    c = canvas.Canvas(translated_file, pagesize=A4)
    total = len(content)

    current_page = 0
    c.showPage()
    batch_size = 32
    for i in range(0, total, batch_size):
        if stop_flag and stop_flag.is_set():
            c.save()
            return "翻译已停止"
        batch = content[i:i + batch_size]
        texts = [item["text"] for item in batch]
        translated_texts = translate_text(texts, src_lang, tgt_lang)
        for idx, item in enumerate(batch):
            if item["page"] != current_page:
                c.showPage()
                current_page = item["page"]
            c.setFont("Helvetica", item["size"])
            c.drawString(item["x"], item["y"], translated_texts[idx])
        if progress_callback:
            progress_callback((i + len(batch)) / total * 100)

    c.save()
    return translated_file


# Translate Word document while preserving format
def translate_word_document(file_path, src_lang, tgt_lang, progress_callback=None, stop_flag=None):
    doc = Document(file_path)
    total_items = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
    items_processed = 0
    batch_size = 32

    paragraphs = [para for para in doc.paragraphs if para.text.strip()]
    for i in range(0, len(paragraphs), batch_size):
        if stop_flag and stop_flag.is_set():
            doc.save(file_path.rsplit(".", 1)[0] + "_partial.docx")
            return "翻译已停止"
        batch = paragraphs[i:i + batch_size]
        texts = [para.text for para in batch]
        translated_texts = translate_text(texts, src_lang, tgt_lang)
        for para, translated in zip(batch, translated_texts):
            para.text = translated
            items_processed += 1
            if progress_callback:
                progress_callback(items_processed / total_items * 100)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if stop_flag and stop_flag.is_set():
                    doc.save(file_path.rsplit(".", 1)[0] + "_partial.docx")
                    return "翻译已停止"
                if cell.text.strip():
                    cell.text = translate_text(cell.text, src_lang, tgt_lang)[0]
                    items_processed += 1
                    if progress_callback:
                        progress_callback(items_processed / total_items * 100)

    translated_file = file_path.rsplit(".", 1)[0] + "_translated.docx"
    doc.save(translated_file)
    return translated_file


# Translate TXT document
def translate_txt_document(file_path, src_lang, tgt_lang, progress_callback=None, stop_flag=None):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    paragraphs = [p for p in text.split("\n") if p.strip()]
    translated_paragraphs = []
    total = len(paragraphs)
    batch_size = 32
    for i in range(0, total, batch_size):
        if stop_flag and stop_flag.is_set():
            translated_file = file_path.rsplit(".", 1)[0] + "_partial.txt"
            with open(translated_file, "w", encoding="utf-8") as f:
                f.write("\n".join(translated_paragraphs))
            return "翻译已停止"
        batch = paragraphs[i:i + batch_size]
        translated_batch = translate_text(batch, src_lang, tgt_lang)
        translated_paragraphs.extend(translated_batch)
        if progress_callback:
            progress_callback((i + len(batch)) / total * 100)

    translated_file = file_path.rsplit(".", 1)[0] + "_translated.txt"
    with open(translated_file, "w", encoding="utf-8") as f:
        f.write("\n".join(translated_paragraphs))
    return translated_file


# Translate entire document
def translate_document(file_path, file_type, src_lang, tgt_lang, progress_callback=None, stop_flag=None):
    if file_type == "pdf":
        return translate_pdf_document(file_path, src_lang, tgt_lang, progress_callback, stop_flag)
    elif file_type == "word":
        return translate_word_document(file_path, src_lang, tgt_lang, progress_callback, stop_flag)
    elif file_type == "txt":
        return translate_txt_document(file_path, src_lang, tgt_lang, progress_callback, stop_flag)
    else:
        return "不支持的文件类型"


# Show translation result in a bubble window
def show_translation_bubble(text, font_size):
    bubble = tk.Toplevel()
    bubble.overrideredirect(True)
    bubble.attributes('-topmost', True)

    screen_width = bubble.winfo_screenwidth()
    mouse_x = bubble.winfo_pointerx()
    max_width = min(900, screen_width - mouse_x - 20)
    label = tk.Label(bubble, text=text, bg="white", padx=10, pady=5, borderwidth=1, relief="solid",
                     font=("Arial", font_size), wraplength=max_width, justify="left")
    label.pack()
    bubble.geometry(f"+{mouse_x}+{bubble.winfo_pointery()}")

    def start_drag(event):
        bubble.x = event.x
        bubble.y = event.y

    def drag(event):
        deltax = event.x - bubble.x
        deltay = event.y - bubble.y
        x = bubble.winfo_x() + deltax
        y = bubble.winfo_y() + deltay
        bubble.geometry(f"+{x}+{y}")

    label.bind("<Button-1>", start_drag)
    label.bind("<B1-Motion>", drag)

    menu = Menu(bubble, tearoff=0)
    menu.add_command(label="销毁", command=bubble.destroy)

    def popup(event):
        menu.post(event.x_root, event.y_root)

    bubble.bind("<Button-3>", popup)


# GUI Application
class TranslationApp:
    def __init__(self, root):
        self.root = root
        self.root.title("本地AI翻译软件")
        self.clipboard_content = ""
        self.is_translating = False
        self.stop_file_translation = threading.Event()
        self.bubble_font_size = 14

        # Language selection
        self.src_lang_label = tk.Label(root, text="源语言:")
        self.src_lang_label.pack(pady=5)
        self.src_lang_var = tk.StringVar(value="自动检测")
        self.src_lang_menu = tk.OptionMenu(root, self.src_lang_var, *languages.keys())
        self.src_lang_menu.pack(pady=5)

        self.tgt_lang_label = tk.Label(root, text="目标语言:")
        self.tgt_lang_label.pack(pady=5)
        self.tgt_lang_var = tk.StringVar(value="中文")
        self.tgt_lang_menu = tk.OptionMenu(root, self.tgt_lang_var, *languages.keys())
        self.tgt_lang_menu.pack(pady=5)

        # Start translation button
        self.start_button = tk.Button(root, text="开启翻译", command=self.start_translation)
        self.start_button.pack(pady=10)

        # Stop translation button
        self.stop_button = tk.Button(root, text="关闭翻译", command=self.stop_translation)
        self.stop_button.pack(pady=10)

        # Upload document button
        self.upload_button = tk.Button(root, text="上传文档", command=self.upload_document)
        self.upload_button.pack(pady=10)

        # Stop file translation button
        self.stop_file_button = tk.Button(root, text="停止文件翻译", command=self.stop_file_translation_func)
        self.stop_file_button.pack(pady=10)

        # Translation result display
        self.result_label = tk.Label(root, text="翻译结果将显示在这里")
        self.result_label.pack(pady=10)

        # Progress bar
        self.progress = Progressbar(root, orient="horizontal", length=200, mode="determinate")
        self.progress.pack(pady=10)

        # Adjust bubble font size
        self.font_size_frame = tk.Frame(root)
        self.font_size_frame.pack(pady=5)
        self.font_size_label = tk.Label(self.font_size_frame, text="气泡字体大小:")
        self.font_size_label.pack(side=tk.LEFT)
        self.font_size_display = tk.Label(self.font_size_frame, text=str(self.bubble_font_size))
        self.font_size_display.pack(side=tk.LEFT, padx=5)
        self.font_size_scale = Scale(self.font_size_frame, from_=10, to=30, orient="horizontal",
                                     command=self.update_font_size)
        self.font_size_scale.set(self.bubble_font_size)
        self.font_size_scale.pack(side=tk.LEFT)

        # Translation state label
        self.translation_state_label = tk.Label(root, text="翻译状态: 关闭")
        self.translation_state_label.pack(pady=5)

        # Create system tray icon
        self.create_tray_icon()

    def create_tray_icon(self):
        # Create an image for the tray icon
        image = Image.new('RGB', (64, 64), color='white')
        draw = ImageDraw.Draw(image)
        draw.rectangle((0, 0, 64, 64), fill='blue')

        # Create the tray icon
        self.tray_icon = pystray.Icon("translation_app", image, "本地AI翻译软件", self.create_tray_menu())
        threading.Thread(target=self.tray_icon.run, daemon=True).start()

    def create_tray_menu(self):
        return pystray.Menu(
            pystray.MenuItem(
                "开启翻译" if not self.is_translating else "关闭翻译",
                self.toggle_translation,
                checked=lambda item: self.is_translating
            ),
            pystray.MenuItem(
                "开启文档翻译" if not self.stop_file_translation.is_set() else "关闭文档翻译",
                self.toggle_file_translation,
                checked=lambda item: not self.stop_file_translation.is_set()
            ),
            pystray.MenuItem("退出", self.exit_app)
        )

    def toggle_translation(self):
        if self.is_translating:
            self.stop_translation()
        else:
            self.start_translation()
        self.tray_icon.update_menu()
        self.update_translation_state_label()

    def toggle_file_translation(self):
        if self.stop_file_translation.is_set():
            self.stop_file_translation.clear()
        else:
            self.stop_file_translation.set()
        self.tray_icon.update_menu()

    def exit_app(self):
        self.tray_icon.stop()
        self.root.quit()

    def update_font_size(self, value):
        self.bubble_font_size = int(float(value))
        self.font_size_display.config(text=str(self.bubble_font_size))

    def start_translation(self):
        if not self.is_translating:
            self.is_translating = True
            threading.Thread(target=self.monitor_clipboard, daemon=True).start()
            messagebox.showinfo("提示", "翻译功能已开启，请复制文字")
            self.update_translation_state_label()

    def stop_translation(self):
        self.is_translating = False
        messagebox.showinfo("提示", "翻译功能已关闭")
        self.update_translation_state_label()

    def update_translation_state_label(self):
        state_text = "开启" if self.is_translating else "关闭"
        self.translation_state_label.config(text=f"翻译状态: {state_text}")

    def monitor_clipboard(self):
        recent_value = pyperclip.paste()
        while self.is_translating:
            current_value = pyperclip.paste()
            if current_value != recent_value and current_value.strip():
                recent_value = current_value
                src_lang = languages[self.src_lang_var.get()]
                tgt_lang = languages[self.tgt_lang_var.get()]
                translated = translate_text(recent_value, src_lang, tgt_lang)[0]
                self.root.after(0, show_translation_bubble, translated, self.bubble_font_size)
            time.sleep(0.2)

    def upload_document(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf"),
                                                          ("Word files", "*.docx"),
                                                          ("Text files", "*.txt")])
        if file_path:
            self.stop_file_translation.clear()
            file_type = "pdf" if file_path.endswith(".pdf") else "word" if file_path.endswith(".docx") else "txt"
            threading.Thread(target=self.translate_and_show, args=(file_path, file_type)).start()

    def stop_file_translation_func(self):
        self.stop_file_translation.set()
        messagebox.showinfo("提示", "文件翻译已停止")

    def translate_and_show(self, file_path, file_type):
        self.progress["value"] = 0
        self.progress["maximum"] = 100

        def update_progress(value):
            self.progress["value"] = value
            self.root.update_idletasks()

        src_lang = languages[self.src_lang_var.get()]
        tgt_lang = languages[self.tgt_lang_var.get()]
        result = translate_document(file_path, file_type, src_lang, tgt_lang, update_progress,
                                    self.stop_file_translation)
        if result == "翻译已停止":
            messagebox.showinfo("提示", "文件翻译已中止，部分翻译结果已保存")
        else:
            messagebox.showinfo("完成", f"翻译完成，文件保存在: {result}")


# Run GUI
if __name__ == "__main__":
    root = tk.Tk()
    app = TranslationApp(root)
    root.mainloop()